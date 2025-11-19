"""
Document processing service for text extraction from various formats.
"""

import os
import tempfile
from pathlib import Path
from typing import Optional

from app.services.video_processor import video_processor


class DocumentProcessor:
    """Process documents (PDF, DOCX, TXT, images) for text extraction."""

    async def extract_text(
        self,
        file_path: str | Path,
        file_type: Optional[str] = None,
    ) -> str:
        """
        Extract text from a document.

        Args:
            file_path: Path to the document
            file_type: Optional file type override

        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Determine file type
        if file_type is None:
            file_type = file_path.suffix.lower().lstrip(".")

        # Route to appropriate extractor
        extractors = {
            "pdf": self._extract_from_pdf,
            "docx": self._extract_from_docx,
            "doc": self._extract_from_docx,
            "txt": self._extract_from_text,
            "md": self._extract_from_text,
            "png": self._extract_from_image,
            "jpg": self._extract_from_image,
            "jpeg": self._extract_from_image,
            "gif": self._extract_from_image,
            "bmp": self._extract_from_image,
            "tiff": self._extract_from_image,
        }

        extractor = extractors.get(file_type)
        if extractor is None:
            raise ValueError(f"Unsupported file type: {file_type}")

        return await extractor(file_path)

    async def extract_text_from_bytes(
        self,
        data: bytes,
        file_type: str,
    ) -> str:
        """Extract text from document bytes."""
        suffix = f".{file_type}"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name

        try:
            return await self.extract_text(tmp_path, file_type)
        finally:
            os.unlink(tmp_path)

    async def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file."""
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n\n".join(text_parts)

        # If PDF has no extractable text, try OCR
        if not text.strip():
            text = await self._ocr_pdf(file_path)

        return text

    async def _ocr_pdf(self, file_path: Path) -> str:
        """OCR a PDF file by converting pages to images."""
        try:
            import subprocess

            # Convert PDF to images using pdftoppm if available
            with tempfile.TemporaryDirectory() as tmp_dir:
                result = subprocess.run(
                    [
                        "pdftoppm", "-png",
                        str(file_path),
                        f"{tmp_dir}/page"
                    ],
                    capture_output=True,
                    timeout=60,
                )

                if result.returncode != 0:
                    return ""

                # OCR each page image
                text_parts = []
                for img_file in sorted(Path(tmp_dir).glob("*.png")):
                    ocr_result = await video_processor.extract_text_from_image(img_file)
                    if ocr_result.text:
                        text_parts.append(ocr_result.text)

                return "\n\n".join(text_parts)
        except Exception:
            return ""

    async def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file."""
        from docx import Document

        doc = Document(str(file_path))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    async def _extract_from_text(self, file_path: Path) -> str:
        """Extract text from a plain text file."""
        return file_path.read_text(encoding="utf-8", errors="ignore")

    async def _extract_from_image(self, file_path: Path) -> str:
        """Extract text from an image using OCR."""
        ocr_result = await video_processor.extract_text_from_image(file_path)
        return ocr_result.text

    def get_supported_formats(self) -> list[str]:
        """Get list of supported document formats."""
        return [
            "pdf", "docx", "doc", "txt", "md",
            "png", "jpg", "jpeg", "gif", "bmp", "tiff"
        ]


# Global instance
document_processor = DocumentProcessor()
